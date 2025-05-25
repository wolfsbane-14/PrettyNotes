import gradio as gr
# from new_v4 import GeminiOutlineConverter
from new_v4 import GeminiContentPreservingConverter # Changed this line
import os
import glob
from dotenv import load_dotenv
import io
import sys
from contextlib import redirect_stdout, redirect_stderr

load_dotenv()
OUTPUT_FOLDER = "generated_docs"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
MANUALLY_ENTERED_API_KEY = None

# Clean old files
for f in glob.glob(f"{OUTPUT_FOLDER}/*.docx"):
    os.remove(f)

# OLD FUNCTION WITH LIMITED FUNCTINALITY
# def convert_pdf_to_outline_simplified(pdf_path):
#     # global MANUALLY_ENTERED_API_KEY

#     if not pdf_path:
#         return "❌ No PDF file provided.", None

#     current_api_key = os.getenv("GEMINI_API_KEY", "").strip()
#     # current_api_key = (api_key_input_from_ui or MANUALLY_ENTERED_API_KEY or os.getenv("GEMINI_API_KEY", "")).strip()

#     # if not current_api_key:
#     #     return "🔐 Gemini API key not found. Please enter it above.", None

#     if not current_api_key:
#         return "🔐 Gemini API key not found in environment variables.", None

#     try:
#         # converter = GeminiOutlineConverter(api_key=current_api_key)
#         converter = GeminiContentPreservingConverter(api_key=current_api_key)
#         # MANUALLY_ENTERED_API_KEY = current_api_key
#     except Exception as e:
#         # MANUALLY_ENTERED_API_KEY = None
#         return f"❌ Failed to initialize Gemini client: {e}", None

#     base_name = os.path.splitext(os.path.basename(pdf_path))[0]
#     output_path = os.path.join(OUTPUT_FOLDER, f"{base_name}_styled_outline.docx")

#     try:
#         result_path = converter.process_file(pdf_path, output_path)
#     except Exception as e:
#         return f"⚠️ Error during processing: {e}", None

#     if result_path and os.path.exists(result_path):
#         try:
#             from docx import Document
#             doc = Document(result_path)
#             if len(doc.paragraphs) < 5 and any("no outlines" in p.text.lower() for p in doc.paragraphs):
#                 return "📄 Processed, but no meaningful outline generated.", result_path
#         except:
#             pass
#         return "✅ Successfully converted!", result_path

#     return "❌ Failed to generate output file.", None

# updated function with more functionality (detailed status log)
def convert_pdf_to_outline_simplified(pdf_path):
    if not pdf_path:
        return "❌ No PDF file provided.", None

    current_api_key = os.getenv("GEMINI_API_KEY", "").strip()

    if not current_api_key:
        return "🔐 Gemini API key not found in environment variables.", None

    # Capture all print outputs
    stdout_capture = io.StringIO()
    stderr_capture = io.StringIO()
    
    try:
        with redirect_stdout(stdout_capture), redirect_stderr(stderr_capture):
            converter = GeminiContentPreservingConverter(api_key=current_api_key)
            
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            output_path = os.path.join(OUTPUT_FOLDER, f"{base_name}_styled_outline.docx")
            
            result_path = converter.process_file(pdf_path, output_path)
        
        # Get captured outputs
        stdout_output = stdout_capture.getvalue()
        stderr_output = stderr_capture.getvalue()
        
        # Filter to show only content preservation messages
        preservation_logs = ""
        if stdout_output:
            lines = stdout_output.split('\n')
            preservation_lines = [line for line in lines if 
                                'Content preservation check:' in line or 
                                'outline passed content preservation check' in line or
                                'outline FAILED strict content preservation check' in line]
            if preservation_lines:
                preservation_logs += f"📊 Content Preservation Results:\n" + "\n".join(preservation_lines) + "\n"
        
        if stderr_output:
            preservation_logs += f"⚠️ Errors:\n{stderr_output}\n"
            
    except Exception as e:
        return f"❌ Failed to initialize Gemini client: {e}", None

    if result_path and os.path.exists(result_path):
        try:
            from docx import Document
            doc = Document(result_path)
            if len(doc.paragraphs) < 5 and any("no outlines" in p.text.lower() for p in doc.paragraphs):
                status_message = f"📄 Processed, but no meaningful outline generated.\n\n{preservation_logs}"
                return status_message, result_path
        except:
            pass
        
        status_message = f"✅ Successfully converted!\n\n{preservation_logs}"
        return status_message, result_path

    status_message = f"❌ Failed to generate output file.\n\n{preservation_logs}"
    return status_message, None

# 🌟 Enhanced Gradio UI with gr.Blocks
# with gr.Blocks(title="PrettyNotes ✨") as demo:
#     gr.Markdown("""
#     # ✨ PrettyNotes
#     Transform any PDF into a **structured, editable DOCX outline** with the power of Gemini AI.

#     ✅ Ideal for lecture notes, research papers, or summaries  
#     🔐 Your API key is never stored  
#     📎 DOCX is fully stylized and ready to edit
#     """)

#     with gr.Row():
#         with gr.Column(scale=3):
#             pdf_input = gr.File(label="📄 Upload Your PDF", type="filepath")
#             api_key_input = gr.Textbox(
#                 label="🔐 Gemini API Key",
#                 placeholder="Chaabi kaha hai??????",
#                 type="password"
#             )
#             convert_button = gr.Button("🚀 Convert to Outline")

#         with gr.Column(scale=2):
#             status_output = gr.Textbox(label="📣 Status", interactive=False)
#             docx_output = gr.File(label="📥 Download DOCX")

#     convert_button.click(
#         convert_pdf_to_outline_simplified,
#         inputs=[pdf_input, api_key_input],
#         outputs=[status_output, docx_output]
#     )

# demo.launch()


# updated code
# with gr.Blocks(css="""
# body {
#     background-image: url('https://images.pexels.com/photos/7135037/pexels-photo-7135037.jpeg?auto=compress&cs=tinysrgb&w=1260&h=750&dpr=1');
#     background-size: cover;
#     background-position: center;
#     background-attachment: fixed;
#     font-family: 'Segoe UI', sans-serif;
# }
# # .gradio-container {
# #     background-color: rgba(255, 255, 255, 0.85);
# #     border-radius: 12px;
# #     padding: 20px;
# #     max-width: 900px;
# #     margin: auto;
# #     box-shadow: 0 0 40px rgba(0, 0, 0, 0.3);
# # }

# # updated code with forced color scheme
# with gr.Blocks(css="""
# body {
#   background-image: url('https://images.pexels.com/photos/7135037/pexels-photo-7135037.jpeg?auto=compress&cs=tinysrgb&w=1260&h=750&dpr=1');
#   background-size: cover;
#   background-position: center;
#   background-attachment: fixed;
#   font-family: 'Segoe UI', sans-serif;
# }

# body, body * {
#   color: white; /* Removed !important */
# }

# .gradio-container,
# .gradio-container *,
# .gr-markdown,
# .gr-markdown *,
# .gr-button,
# .gr-textbox,
# .gr-file,
# label,
# h1, h2, h3, h4, h5, h6,
# p, span, div {
#   color: white;
# }

# @media (prefers-color-scheme: dark) {
#   body, body * {
#     color: white !important;
#   }
# }

# @media (prefers-color-scheme: light) {
#   .gr-button,
#   .gr-textbox,
#   .gr-file,
#   input,
#   button,
#   textarea,
#   label,
#   h1, h2, h3, h4, h5, h6,
#   p, span, div {
#     color: black !important;
#   }
# }
# """, title="PrettyNotes ✨") as demo:

#     gr.Markdown("""
#     # ✨ PrettyNotes
#     Transform any PDF into a **structured, editable DOCX outline** in seconds.
    
#     ✅ Ideal for **lecture notes**, **research papers**, project reports, or study material  
#     🗑️ **Uploaded PDFs and DOCX outputs are automatically deleted** after your session ends  
#     📎 DOCX is **stylized, formatted**, and ready to edit in Word or Google Docs  
#     🌈 Keywords are **color-highlighted** to boost clarity and readability  
#     📄 Works best with text-based PDFs (not scanned image PDFs)
    
#     ---
#     ✨ *Use PrettyNotes to convert chaos into clarity – perfect for students, educators, and lifelong learners!*
#     """)



#     with gr.Row():
#         with gr.Column(scale=3):
#             pdf_input = gr.File(label="📄 Upload Your PDF", type="filepath")
#             # api_key_input = gr.Textbox(
#             #     label="🔐 Gemini API Key",
#             #     placeholder="Enter here",
#             #     type="password"
#             # )
#             convert_button = gr.Button("🚀 Convert to Outline")

#         with gr.Column(scale=2):
#             status_output = gr.Textbox(label="📣 Status", interactive=False)
#             docx_output = gr.File(label="📥 Download DOCX")

#     convert_button.click(
#         convert_pdf_to_outline_simplified,
#         inputs=[pdf_input],
#         outputs=[status_output, docx_output]
#     )

# demo.launch()

# updated code with forced color and tweaks for white mode
with gr.Blocks(
    css="""
    body {
      background-image: url('https://images.pexels.com/photos/7135037/pexels-photo-7135037.jpeg?auto=compress&cs=tinysrgb&w=1260&h=750&dpr=1');
      background-size: cover;
      background-position: center;
      background-attachment: fixed;
      font-family: 'Segoe UI', sans-serif;
    }
    body, body * {
      color: white;
    }
    .gradio-container,
    .gradio-container *,
    .gr-markdown,
    .gr-markdown *,
    .gr-button,
    .gr-textbox,
    .gr-file,
    label,
    h1, h2, h3, h4, h5, h6,
    p, span, div {
      color: white;
    }
    @media (prefers-color-scheme: dark) {
      body, body * {
        color: white !important;
      }
    }
    @media (prefers-color-scheme: light) {
      .gr-button,
      .gr-textbox,
      .gr-file,
      input,
      button,
      textarea,
      label,
      h1, h2, h3, h4, h5, h6,
      p, span, div {
        color: black !important;
      }
      #top-markdown {
        background-color: rgba(255, 255, 255, 0.8);
        border-radius: 10px;
        padding: 10px;
      }
      .app-row, .app-column {
        # border: 1px solid lightgray;
        background-color: rgba(255, 255, 255, 0.8);
        border-radius: 10px;
        padding: 10px;
      }
    }
    """,
    title="PrettyNotes ✨"
) as demo:
    gr.Markdown("""
    # ✨ PrettyNotes
    Transform any PDF into a **structured, editable DOCX outline** in seconds.
    
    ✅ Ideal for **lecture notes**, **research papers**, project reports, or study material  
    🗑️ **Uploaded PDFs and DOCX outputs are automatically deleted** after your session ends  
    📎 DOCX is **stylized, formatted**, and ready to edit in Word or Google Docs  
    🌈 Keywords are **color-highlighted** to boost clarity and readability  
    📄 Works best with text-based PDFs (not scanned image PDFs)
    
    ---
    ✨ *Use PrettyNotes to convert chaos into clarity – perfect for students, educators, and lifelong learners!*
    """, elem_id="top-markdown")

    # FOR FUTURE UPDATES
    # model_selector = gr.Radio(choices=["Gemini", "DeepSeek", "LLaMA", "Mistral"], label="Select Model")

    # # Font settings
    # heading_font_dropdown = gr.Dropdown(choices=["Courier New", "Arial", "Times New Roman", "Calibri"], label="Heading Font", value="Courier New")
    # heading_font_size = gr.Slider(10, 24, value=14, step=1, label="Heading Font Size")
    # body_font_size = gr.Slider(8, 16, value=12, step=1, label="Body Font Size")

    with gr.Row(elem_classes="app-row"):
        with gr.Column(scale=3, elem_classes="app-column"):
            pdf_input = gr.File(label="📄 Upload Your PDF", type="filepath")
            convert_button = gr.Button("🚀 Convert to Outline")
        with gr.Column(scale=2, elem_classes="app-column"):
            status_output = gr.Textbox(label="📣 Status", interactive=False)
            docx_output = gr.File(label="📥 Download DOCX")

    convert_button.click(
        convert_pdf_to_outline_simplified,
        inputs=[pdf_input],
        outputs=[status_output, docx_output]
    )

demo.launch()

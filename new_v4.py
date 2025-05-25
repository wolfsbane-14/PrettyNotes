import os
import re
import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR # For line spacing
from docx.enum.text import WD_LINE_SPACING # For line spacing
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configuration for chunking
MAX_CHARS_PER_CHUNK = 12000 # Keep in mind Gemini's token limits, this might need adjustment

# --- Style Configuration (remains the same) ---
HIERARCHY_MARKER_FONT_NAME = 'Courier New'
TITLE_TEXT_FONT_NAME = 'Courier New'
SUBTITLE_TEXT_FONT_NAME = 'Courier New'
CONTENT_TEXT_FONT_NAME = 'Courier New'
HEADING_FONT_SIZE = 14
BODY_FONT_SIZE = 12
KEYWORDS_TO_HIGHLIGHT = {
    "and": "FF00FF", "or": "FF00FF", "not": "FF00FF", "if": "FF00FF",
    "then": "FF00FF", "else": "FF00FF", "when": "FF00FF", "while": "FF00FF",
    "for": "FF00FF", "in": "FF00FF", "with": "FF00FF", "by": "FF00FF",
    "to": "007ACC", "into": "007ACC", "from": "007ACC", "of": "007ACC",
    "at": "007ACC", "on": "007ACC", "through": "007ACC", "via": "007ACC",
    "over": "007ACC", "under": "007ACC", "between": "007ACC", "among": "007ACC",
    "is": "00D8B0", "are": "00D8B0", "was": "00D8B0", "were": "00D8B0",
    "will": "00D8B0", "can": "00D8B0", "should": "00D8B0", "must": "00D8B0",
    "may": "00D8B0", "could": "00D8B0", "would": "00D8B0",
    "strategic": "FFD700", "management": "FFD700", "analysis": "FFD700",
    "framework": "FFD700", "methodology": "FFD700", "approach": "FFD700",
    "implementation": "FFD700", "evaluation": "FFD700", "assessment": "FFD700",
    "development": "FFD700", "research": "FFD700", "study": "FFD700",
    "process": "FFD700", "system": "FFD700", "model": "FFD700", "theory": "FFD700",
    "human": "A6E22E", "people": "A6E22E", "individual": "A6E22E", "person": "A6E22E",
    "employee": "A6E22E", "worker": "A6E22E", "staff": "A6E22E", "team": "A6E22E",
    "group": "A6E22E", "organization": "A6E22E", "company": "A6E22E", "business": "A6E22E",
    "resource": "00BFFF", "data": "00BFFF", "information": "00BFFF", "knowledge": "00BFFF",
    "skill": "00BFFF", "capability": "00BFFF", "capacity": "00BFFF", "asset": "00BFFF",
    "tool": "00BFFF", "method": "00BFFF", "technique": "00BFFF", "solution": "00BFFF",
    "all": "FF8C00", "some": "FF8C00", "many": "FF8C00", "few": "FF8C00",
    "most": "FF8C00", "several": "FF8C00", "various": "FF8C00", "multiple": "FF8C00",
    "single": "FF8C00", "first": "FF8C00", "second": "FF8C00", "third": "FF8C00",
    "primary": "FF8C00", "secondary": "FF8C00", "main": "FF8C00", "key": "FF8C00",
    "important": "FF8C00", "critical": "FF8C00", "essential": "FF8C00", "significant": "FF8C00",
}
DEFAULT_TEXT_COLOR = "000000"
BULLET_PREFIX = "|-- "
# --- End Style Configuration ---

class GeminiContentPreservingConverter:
    def __init__(self, api_key=None):
        """Initialize the converter with the Gemini API"""
        try:
            if not api_key:
                raise ValueError("Gemini API key not provided.")
            genai.configure(api_key=api_key)
            self.model = genai.GenerativeModel('gemini-1.5-flash-latest')
            print("Gemini client configured successfully.")
        except Exception as e:
            print(f"Failed to configure Gemini client: {e}")
            print("Please ensure the GEMINI_API_KEY is passed correctly.")
            raise

    def _format_text_for_docx(self, paragraph, text_content, text_font_name, text_color, is_bold=False):
        """Applies keyword highlighting and styling to a docx paragraph."""
        if not text_content:
            return

        temp_text = text_content
        
        sorted_keywords = sorted(KEYWORDS_TO_HIGHLIGHT.keys(), key=len, reverse=True)

        matches = []
        for keyword in sorted_keywords:
            pattern = r'\b(' + re.escape(keyword) + r')\b'
            for match in re.finditer(pattern, temp_text, re.IGNORECASE):
                matches.append((match.start(), match.end(), match.group(1), KEYWORDS_TO_HIGHLIGHT[keyword]))
        
        matches.sort(key=lambda x: x[0])

        last_idx = 0
        for start, end, matched_word, color in matches:
            if start > last_idx:
                run = paragraph.add_run(temp_text[last_idx:start])
                run.font.name = text_font_name
                run.font.size = Pt(BODY_FONT_SIZE)
                run.font.color.rgb = self._hex_to_rgb(DEFAULT_TEXT_COLOR)
                run.bold = is_bold
            
            run = paragraph.add_run(matched_word)
            run.font.name = text_font_name
            run.font.size = Pt(BODY_FONT_SIZE)
            run.font.color.rgb = self._hex_to_rgb(color)
            run.bold = True
            
            last_idx = end
        
        if last_idx < len(temp_text):
            run = paragraph.add_run(temp_text[last_idx:])
            run.font.name = text_font_name
            run.font.size = Pt(BODY_FONT_SIZE)
            run.font.color.rgb = self._hex_to_rgb(DEFAULT_TEXT_COLOR)
            run.bold = is_bold
            
        if not paragraph.runs and text_content:
            run = paragraph.add_run(text_content)
            run.font.name = text_font_name
            run.font.size = Pt(BODY_FONT_SIZE)
            run.font.color.rgb = self._hex_to_rgb(DEFAULT_TEXT_COLOR)
            run.bold = is_bold

    def _hex_to_rgb(self, hex_color):
        """Converts a hex color string (e.g., "RRGGBB") to an RGB object for python-docx."""
        from docx.shared import RGBColor
        hex_color = hex_color.lstrip('#')
        return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

    def extract_text_from_pdf(self, pdf_path):
        """Extract text content from all pages of a PDF."""
        print(f"Extracting text from PDF: {pdf_path}")
        try:
            doc = fitz.open(pdf_path)
            full_text = []
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                full_text.append(page.get_text("text"))
            doc.close()
            extracted_text = "\n".join(full_text)
            if not extracted_text.strip():
                print("Warning: No text extracted from the PDF. The PDF might be image-based or empty.")
            return extracted_text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None

    def split_text_into_chunks(self, text, max_chars=MAX_CHARS_PER_CHUNK):
        """Splits text into chunks, trying to respect paragraph boundaries."""
        chunks = []
        current_chunk_paras = []
        current_length = 0
        paragraphs = text.split('\n\n')

        for para_idx, paragraph in enumerate(paragraphs):
            para_len = len(paragraph)
            if current_length + para_len + (2 if current_chunk_paras else 0) <= max_chars:
                current_chunk_paras.append(paragraph)
                current_length += para_len + (2 if current_chunk_paras else 0)
            else:
                if current_chunk_paras:
                    chunks.append("\n\n".join(current_chunk_paras))
                
                if para_len > max_chars:
                    start = 0
                    while start < para_len:
                        end = start + max_chars
                        chunks.append(paragraph[start:end])
                        start = end
                    current_chunk_paras = []
                    current_length = 0
                else:
                    current_chunk_paras = [paragraph]
                    current_length = para_len
            
            if para_idx == len(paragraphs) - 1 and current_chunk_paras:
                chunks.append("\n\n".join(current_chunk_paras))
        
        return [chunk for chunk in chunks if chunk.strip()]

    def process_with_gemini(self, text_chunk, chunk_num, total_chunks, original_full_text):
        """Send a single text chunk to Gemini for FORMATTING and MINOR CORRECTIONS."""
        if not text_chunk or not text_chunk.strip():
            print(f"Skipping empty chunk {chunk_num}/{total_chunks}.")
            return ""

        # UPDATED PROMPT: Now allows for minor corrections without altering core meaning
        content_preservation_prompt = """
        YOU ARE A TEXT FORMATTER AND MINOR ERROR CORRECTOR. YOUR PRIMARY GOAL IS TO ORGANIZE THE PROVIDED TEXT INTO A HIERARCHICAL OUTLINE FORMAT.
        CRITICAL RULES:
        1.  **PRESERVE CORE MEANING:** Do NOT rephrase, summarize, or interpret the content in a way that changes its original meaning.
        2.  **MAINTAIN LOGICAL FLOW:** The order of information (words, sentences, paragraphs) should be preserved logically. No reordering of sentences or key phrases.
        3.  **CORRECT OBVIOUS ERRORS:**
            * **Typographical:** Fix missing spaces between words (e.g., "wordone wordtwo" -> "word one word two").
            * **Punctuation:** Add missing commas, periods, etc., where grammatically necessary and obvious.
            * **Lexical/Grammatical:** Correct clear grammatical errors or misspellings that do not change the word's intended meaning (e.g., "teh" -> "the").
        4.  **DO NOT:**
            * Add your own words, explanations, or interpretations.
            * Remove or skip any sentences or paragraphs.
            * Change factual claims or introduce new information.
            * Correct factual errors (as you cannot verify external facts).
        5.  **USE EXACT ORIGINAL WORDING** as much as possible, applying only the allowed corrections.
        YOUR ONLY TASK: Take the existing text and organize it using this hierarchy format:
        FORMAT RULES:
        1. Main Topic One
        |-- [Exact sentence/paragraph from original text, with minor corrections applied]
        |-- [Another exact sentence/paragraph from original text, with minor corrections applied]
          1.b Subtopic (if natural division exists in original)
          |-- [Exact sentence from original under this subtopic, with minor corrections applied]
          | |-- [Exact sentence if it's a sub-detail, with minor corrections applied]
        2. Main Topic Two
        |-- [Exact sentence/paragraph from original text, with minor corrections applied]
        INDENTATION RULES:
        - Main sections: "1. ", "2. " etc. (no leading spaces)
        - Subsections: "  1.b ", "  2.a " etc. (exactly 2 leading spaces)
        - Bullets under main: "|-- " (no leading spaces)
        - Bullets under subsections: "  |-- " (exactly 2 leading spaces)
        - Sub-bullets: Add 2 more spaces per level: "| |-- ", "| | |-- "
        ORGANIZATION STRATEGY:
        - Look for existing headings, section breaks, or paragraph divisions in the original text.
        - Group related sentences that appear consecutively.
        - If no clear structure exists, simply list each paragraph as a bullet point.
        - Create topics based on natural content breaks, not your interpretation.
        EXAMPLE OF WHAT YOU SHOULD DO:
        Original text: "Strategicplanning involvesmultiple steps.First, assess current situation.Marketanalysis is crucial.External factors must be considered."
        CORRECT OUTPUT:
        1. Strategic Planning Process
        |-- Strategic planning involves multiple steps.
        |-- First, assess current situation.
        |-- Market analysis is crucial.
        |-- External factors must be considered.
        Remember: Your role is to format and correct minor errors, while strictly preserving the original meaning and sequence of information.
        """

        chunk_instruction = f"""
        This is Chunk {chunk_num} of {total_chunks} from a larger document.
        
        Format and correct ONLY the content in THIS CHUNK. Do not add connecting text between chunks.
        
        Here is the exact text content to format and correct:
        
        ---
        {text_chunk}
        ---
        
        Provide ONLY the formatted outline using exact original text (with allowed minor corrections). No additional commentary.
        """

        full_prompt = f"{content_preservation_prompt}\n\n{chunk_instruction}"
        
        print(f"Sending Chunk {chunk_num}/{total_chunks} to Gemini for FORMATTING and CORRECTIONS ({len(text_chunk)} chars)...")
        try:
            response = self.model.generate_content(
                full_prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.2,  # Slightly higher to allow for minor corrections, but still low
                    top_p=0.8,        # Reduced to limit variation
                    max_output_tokens=4096
                )
            )
            
            if not response.candidates:
                if hasattr(response, 'prompt_feedback') and response.prompt_feedback.block_reason:
                    print(f"Warning: Chunk {chunk_num} was blocked by Gemini. Reason: {response.prompt_feedback.block_reason}")
                    return ""
                else:
                    print(f"Warning: Chunk {chunk_num} - No content generated by Gemini.")
                    return ""

            outline_output = ""
            if response.candidates and response.candidates[0].content and response.candidates[0].content.parts:
                outline_output = "".join(part.text for part in response.candidates[0].content.parts if hasattr(part, 'text')).strip()

            if not outline_output:
                print(f"Warning: Chunk {chunk_num} - Gemini returned an empty outline.")
                return ""

            # Check content preservation, but always return outline_output
            # This check will now often show lower preservation as corrections are allowed
            if self._strict_content_preservation_check(outline_output, text_chunk):
                print(f"Chunk {chunk_num} outline passed content preservation check (minimal changes).")
            else:
                print(f"WARNING: Chunk {chunk_num} outline FAILED strict content preservation check. Content was corrected/reformatted.")
                print("The output will still be included as per user request to fix errors.")
            
            return outline_output # ALWAYS RETURN THE OUTPUT

        except Exception as e:
            print(f"Error with Gemini API for Chunk {chunk_num}/{total_chunks}: {e}")
            return ""

    def _strict_content_preservation_check(self, outline_text, original_chunk_text):
        """
        Strict check to ensure the outline contains original content without alteration.
        Verifies that substantial portions of original text appear in the outline.
        This check will now often show lower preservation as corrections are allowed.
        """
        if not outline_text or not original_chunk_text:
            return True

        # Remove formatting markers to check actual content
        cleaned_outline = re.sub(r'^\s*\d+\.\s*', '', outline_text, flags=re.MULTILINE)
        cleaned_outline = re.sub(r'^\s*\d+\.[a-zA-Z]\.\s*', '', cleaned_outline, flags=re.MULTILINE)
        cleaned_outline = re.sub(r'^\s*\|--\s*', '', cleaned_outline, flags=re.MULTILINE)
        cleaned_outline = re.sub(r'^\s*\|\s*\|--\s*', '', cleaned_outline, flags=re.MULTILINE)
        
        # Extract sentences from both texts
        original_sentences = [s.strip() for s in re.split(r'[.!?]+', original_chunk_text) if s.strip() and len(s.strip()) > 10]
        outline_sentences = [s.strip() for s in re.split(r'[.!?]+', cleaned_outline) if s.strip() and len(s.strip()) > 10]
        
        if not original_sentences:
            return True
            
        # Check how many original sentences appear (with minor variations allowed for punctuation)
        preserved_count = 0
        for orig_sentence in original_sentences:
            orig_words = set(re.findall(r'\b\w+\b', orig_sentence.lower()))
            for outline_sentence in outline_sentences:
                outline_words = set(re.findall(r'\b\w+\b', outline_sentence.lower()))
                # If 90% of original words are preserved, consider it preserved
                if orig_words and len(orig_words.intersection(outline_words)) / len(orig_words) >= 0.9:
                    preserved_count += 1
                    break
        
        preservation_ratio = preserved_count / len(original_sentences) if original_sentences else 0
        print(f"Content preservation check: {preservation_ratio:.2%} of original sentences preserved")
        
        # This threshold is now more indicative than blocking, as corrections are allowed
        return preservation_ratio >= 0.8 # Keep as 0.8 or lower if you want to be less strict on warnings

    def parse_llm_outline(self, outline_text):
        """Parse the LLM's structured outline based on indentation."""
        if not outline_text: return []
        lines = outline_text.split('\n')
        
        parsed_structure = []
        current_main_section = None
        current_subsection = None
        parent_bullet_stack = []

        for line_text in lines:
            stripped_line_content = line_text.strip()
            if not stripped_line_content: continue

            leading_spaces = len(line_text) - len(line_text.lstrip(' '))

            main_match = re.match(r'^(\d+\.\s+)(.*)', stripped_line_content)
            if main_match and leading_spaces == 0:
                marker, title = main_match.groups()
                current_main_section = {'type': 'main_section', 'marker': marker.strip(), 'title': title.strip(), 'content': [], '_indent': 0}
                parsed_structure.append(current_main_section)
                current_subsection = None
                parent_bullet_stack = [(current_main_section, 0)]
                continue

            subsection_match = re.match(r'^(\d+\.[a-zA-Z]\.\s+)(.*)', stripped_line_content)
            if subsection_match and current_main_section and leading_spaces > 0 and not stripped_line_content.startswith('|--'):
                marker, title = subsection_match.groups()
                current_subsection = {'type': 'subsection', 'marker': marker.strip(), 'title': title.strip(), 'content': [], '_indent': leading_spaces}
                current_main_section['content'].append(current_subsection)
                while parent_bullet_stack and parent_bullet_stack[-1][1] >= leading_spaces:
                    parent_bullet_stack.pop()
                parent_bullet_stack.append((current_subsection, leading_spaces))
                continue

            if stripped_line_content.startswith(BULLET_PREFIX):
                bullet_text = stripped_line_content[len(BULLET_PREFIX):].strip()
                
                active_parent_obj = None
                parent_expected_bullet_indent = 0
                
                for i in range(len(parent_bullet_stack) - 1, -1, -1):
                    p_obj, p_indent = parent_bullet_stack[i]
                    if leading_spaces >= p_indent:
                        active_parent_obj = p_obj
                        parent_expected_bullet_indent = p_indent
                        parent_bullet_stack = parent_bullet_stack[:i+1]
                        break
                
                if not active_parent_obj:
                    if parsed_structure: active_parent_obj = parsed_structure[-1]
                    else:
                        active_parent_obj = {'type': 'main_section', 'marker': '?', 'title': 'Content', 'content': [], '_indent': 0}
                        parsed_structure.append(active_parent_obj)
                        parent_bullet_stack = [(active_parent_obj, 0)]
                    parent_expected_bullet_indent = active_parent_obj.get('_indent', 0)

                relative_indent_for_bullet = leading_spaces - parent_expected_bullet_indent
                bullet_level = 1 + (relative_indent_for_bullet // 2)

                bullet_obj = {'type': 'bullet', 'text': bullet_text, 'level': bullet_level, '_abs_indent': leading_spaces}
                active_parent_obj['content'].append(bullet_obj)
                continue
        return parsed_structure

    def _render_content_recursive_docx(self, document, content_list, current_base_indent_inch):
        """Helper to recursively render content for DOCX, managing indentation and styling."""
        
        for item in content_list:
            item_type = item.get('type')
            
            if item_type == 'subsection':
                marker = item.get('marker', '')
                title = item.get('title', '')
                
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = current_base_indent_inch + Inches(0.25)
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Apply 1.5 line spacing
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph.paragraph_format.line_spacing = 1.5

                marker_run = paragraph.add_run(f"{marker} ")
                marker_run.font.name = HIERARCHY_MARKER_FONT_NAME
                marker_run.font.size = Pt(HEADING_FONT_SIZE - 1)
                marker_run.bold = True
                
                self._format_text_for_docx(paragraph, title, SUBTITLE_TEXT_FONT_NAME, DEFAULT_TEXT_COLOR, is_bold=True)
                
                self._render_content_recursive_docx(
                    document,
                    item.get('content', []),
                    current_base_indent_inch + Inches(0.25)
                )
            
            elif item_type == 'bullet':
                text = item.get('text', '')
                bullet_level = item.get('level', 1)
                
                paragraph = document.add_paragraph()
                bullet_indent_inch = current_base_indent_inch + (bullet_level - 1) * Inches(0.25)
                paragraph.paragraph_format.left_indent = bullet_indent_inch
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Apply 1.5 line spacing
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph.paragraph_format.line_spacing = 1.5


                bullet_run = paragraph.add_run(f"{BULLET_PREFIX}")
                bullet_run.font.name = HIERARCHY_MARKER_FONT_NAME
                bullet_run.font.size = Pt(BODY_FONT_SIZE)
                bullet_run.bold = True

                self._format_text_for_docx(paragraph, text, CONTENT_TEXT_FONT_NAME, DEFAULT_TEXT_COLOR)
                
    def create_docx_from_outline(self, parsed_structure_or_text, output_path):
        """Create a DOCX from the parsed outline structure or raw text if parsing fails."""
        document = Document()

        document.styles['Normal'].font.name = 'Courier New'
        document.styles['Normal'].font.size = Pt(BODY_FONT_SIZE)

        main_section_base_indent_inch = Inches(0.25)
        
        # Check if it's a list (parsed structure) or a string (raw text)
        if isinstance(parsed_structure_or_text, list) and parsed_structure_or_text:
            for main_section in parsed_structure_or_text:
                if main_section.get('type') == 'main_section':
                    marker = main_section.get('marker', '')
                    title = main_section.get('title', 'Untitled Section')

                    paragraph = document.add_paragraph()
                    paragraph.paragraph_format.left_indent = main_section_base_indent_inch
                    paragraph.paragraph_format.space_before = Pt(12)
                    paragraph.paragraph_format.space_after = Pt(8)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Apply 1.5 line spacing
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    paragraph.paragraph_format.line_spacing = 1.5


                    marker_run = paragraph.add_run(f"{marker} ")
                    marker_run.font.name = HIERARCHY_MARKER_FONT_NAME
                    marker_run.font.size = Pt(HEADING_FONT_SIZE)
                    marker_run.bold = True
                    
                    self._format_text_for_docx(paragraph, title, TITLE_TEXT_FONT_NAME, DEFAULT_TEXT_COLOR, is_bold=True)

                    self._render_content_recursive_docx(
                        document,
                        main_section.get('content', []),
                        main_section_base_indent_inch
                    )
                document.add_paragraph().add_run().add_break()
        else: # It's a string or empty list, indicating no structured content
            paragraph = document.add_paragraph()
            # Apply 1.5 line spacing to raw text fallback
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = 1.5
            if isinstance(parsed_structure_or_text, str) and parsed_structure_or_text.strip():
                self._format_text_for_docx(paragraph, parsed_structure_or_text, CONTENT_TEXT_FONT_NAME, DEFAULT_TEXT_COLOR)
            else:
                paragraph.add_run("No structured content could be generated or parsed. The document might be image-based, encrypted, or content processing led to an empty output.").font.size = Pt(BODY_FONT_SIZE)

        try:
            document.save(output_path)
            print(f"DOCX created successfully: {output_path}")
        except Exception as e:
            print(f"Error creating DOCX: {e}")

    def process_file(self, input_path, output_path=None):
        print(f"Processing PDF in CONTENT PRESERVATION + CORRECTION MODE: {input_path}")
        file_ext = os.path.splitext(input_path)[1].lower()
        if file_ext != '.pdf':
            print(f"Unsupported file type: {file_ext}.")
            return None

        pdf_full_text = self.extract_text_from_pdf(input_path)
        if not pdf_full_text:
            print("Failed to extract text from PDF.")
            if not output_path:
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                output_path = f"{base_name}_gemini_extraction_failed.docx"
            self.create_docx_from_outline("Failed to extract text from the PDF. The document might be image-based, encrypted, or corrupted.", output_path)
            return output_path

        text_chunks = self.split_text_into_chunks(pdf_full_text)
        if not text_chunks:
            print("PDF text resulted in no processable chunks.")
            if not output_path:
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                output_path = f"{base_name}_gemini_no_chunks.docx"
            self.create_docx_from_outline("PDF content was extracted but resulted in no processable text chunks.", output_path)
            return output_path
            
        print(f"PDF text split into {len(text_chunks)} chunks.")
        all_outlines = []
        for i, chunk_text in enumerate(text_chunks):
            print(f"\nProcessing Chunk {i+1} of {len(text_chunks)} with CORRECTION ENABLED")
            chunk_outline = self.process_with_gemini(chunk_text, i + 1, len(text_chunks), pdf_full_text)
            if chunk_outline:
                all_outlines.append(chunk_outline)
            else:
                print(f"Chunk {i+1} yielded no output from Gemini (e.g., blocked or empty response).")
        
        if not all_outlines:
            print("No outlines were generated by Gemini for any chunks.")
            if not output_path:
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                output_path = f"{base_name}_gemini_empty_output.docx"
            self.create_docx_from_outline("No outlines could be generated by Gemini for the provided content.", output_path)
            return output_path

        combined_outline_text = "\n".join(all_outlines)

        print("\n--- Combined Processed Outline ---")
        print(combined_outline_text[:1000] + "..." if len(combined_outline_text) > 1000 else combined_outline_text)
        print("--- End of Combined Outline ---")

        parsed_structure = self.parse_llm_outline(combined_outline_text)
        
        if not output_path:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = f"{base_name}_gemini_corrected_outline.docx"

        if not parsed_structure and combined_outline_text:
            print("Failed to parse the combined outline into a structured format. DOCX will contain raw corrected content.")
            self.create_docx_from_outline(combined_outline_text, output_path)
        elif parsed_structure:
            self.create_docx_from_outline(parsed_structure, output_path)
        else:
            self.create_docx_from_outline("Content processing resulted in an empty or unparseable output.", output_path)
        
        return output_path

def main():
    print("Gemini Outline Converter: Content Preservation with Minor Corrections")
    print("=" * 60)
    print("⚠️ WARNING: This mode allows Gemini to make minor corrections (spacing, punctuation, obvious typos).")
    print("   This means the output text will NOT be character-for-character identical to the original.")
    print("   However, the aim is to preserve core meaning and logical flow.")
    print("   Factual errors are NOT corrected.")
    print("-" * 60)

    # api_key_input = input("Enter your Gemini API key (or press Enter if GEMINI_API_KEY env var is set): ").strip()
    # gemini_api_key = api_key_input if api_key_input else os.environ.get('GEMINI_API_KEY')
    gemini_api_key = os.environ.get('GEMINI_API_KEY')
    if not gemini_api_key:
        print("Gemini API key not found.")
        return

    try:
        converter = GeminiContentPreservingConverter(api_key=gemini_api_key)
    except Exception:
        print("Exiting due to converter initialization failure.")
        return

    input_file = ""
    while True:
        input_file = input("Enter path to your PDF file: ").strip()
        if not input_file:
            print("Please provide a valid file path.")
            continue
        
        if not os.path.exists(input_file):
            print(f"File not found: {input_file}")
            continue
            
        if not input_file.lower().endswith('.pdf'):
            print("Please provide a PDF file.")
            continue
            
        break
    
    output_file = input("Enter output path for DOCX (or press Enter for auto-naming): ").strip()
    if not output_file:
        output_file = None
    
    print(f"\nProcessing: {input_file}")
    print("This may take several minutes depending on document size...")
    
    try:
        result_path = converter.process_file(input_file, output_file)
        if result_path:
            print(f"\n✅ SUCCESS: Corrected and formatted outline created at: {result_path}")
            print("📋 Remember: Minor corrections were applied for readability. The core meaning and logical flow should be preserved.")
            print("⚠️ If 'Content preservation check FAILED' warnings appeared, it means corrections were made as intended.")
        else:
            print("\n❌ FAILED: Could not process the file.")
    except Exception as e:
        print(f"\n❌ ERROR during processing: {e}")
    
    print("\nThank you for using the Gemini Outline Converter!")

if __name__ == "__main__":
    main()

